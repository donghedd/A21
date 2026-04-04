"""
Word Document Loader
"""
from typing import List
from .base import BaseLoader, Document


class WordLoader(BaseLoader):
    """Load Word documents using python-docx"""
    
    def load(self, file_path: str) -> List[Document]:
        """Load Word file and extract text"""
        try:
            from docx import Document as DocxDocument
            
            doc = DocxDocument(file_path)
            
            # Extract all paragraphs
            paragraphs = []
            for para in doc.paragraphs:
                if para.text.strip():
                    paragraphs.append(para.text)
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        paragraphs.append(' | '.join(row_text))
            
            content = '\n\n'.join(paragraphs)
            
            if content.strip():
                return [Document(
                    page_content=content,
                    metadata={
                        'source': file_path,
                        'file_type': 'docx'
                    }
                )]
            return []
            
        except Exception as e:
            raise Exception(f"Failed to load Word file: {str(e)}")
